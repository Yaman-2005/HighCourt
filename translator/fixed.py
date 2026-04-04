import os
import csv
import torch
import re
import fitz
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from transformers import AutoModelForSeq2SeqLM, AutoTokenizer
from IndicTransToolkit.processor import IndicProcessor

# ---------------------------
# ENV
# ---------------------------
os.environ['HF_HOME'] = r"D:\HC\hf_cache"
os.environ['TEMP'] = r"D:\HC\temp_env"

# ---------------------------
# LOAD GLOSSARY
# ---------------------------
def load_glossary(csv_path):
    glossary = {}
    if not os.path.exists(csv_path):
        return glossary

    with open(csv_path, encoding='utf-8-sig') as f:
        reader = csv.reader(f)
        next(reader, None)
        for row in reader:
            if len(row) >= 2:
                glossary[row[0].strip()] = row[1].strip()

    return glossary

# ---------------------------
# APPLY GLOSSARY (POST ONLY)
# ---------------------------
def apply_glossary(text, glossary):
    for eng, ben in glossary.items():
        text = text.replace(ben, ben)
    return text

# ---------------------------
# CLEAN PAGE INTO LINES
# ---------------------------
def clean_page_lines(text):
    lines = text.split('\n')
    cleaned = []

    for line in lines:
        line = line.strip()

        if not line:
            cleaned.append("")
            continue

        # Remove headers / junk
        if re.search(r'IN THE HIGH COURT|SUPREME COURT|S\.C\.R\.|REPORTS', line, re.I):
            continue

        if re.match(r'^\d+$', line):
            continue

        if line in list("ABCDEFGH"):
            continue

        line = re.sub(r'<.*?>', '', line)

        cleaned.append(line)

    return cleaned

# ---------------------------
# GROUP INTO PARAGRAPHS
# ---------------------------
def group_paragraphs(lines):
    paragraphs = []
    current = ""

    for line in lines:
        if line == "":
            if current:
                paragraphs.append(current.strip())
                current = ""
        else:
            if current.endswith(('.', ':', '?')):
                paragraphs.append(current.strip())
                current = line
            else:
                current += " " + line

    if current:
        paragraphs.append(current.strip())

    return paragraphs

# ---------------------------
# TITLE DETECTION
# ---------------------------
def is_title(text):
    return (
        len(text.split()) <= 12 and
        (text.isupper() or ':' in text or 'Vs.' in text)
    )

# ---------------------------
# CHUNKING
# ---------------------------
def split_into_chunks(text, max_tokens=180):
    sentences = re.split(r'(?<=[.!?])\s+', text)

    chunks = []
    current = ""

    for s in sentences:
        if len((current + " " + s).split()) < max_tokens:
            current += " " + s
        else:
            chunks.append(current.strip())
            current = s

    if current:
        chunks.append(current.strip())

    return chunks

# ---------------------------
# CLEAN OUTPUT
# ---------------------------
def clean_output(text):
    text = re.sub(r'\b(\w+)( \1\b)+', r'\1', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

# ---------------------------
# NORMALIZE BENGALI
# ---------------------------
def normalize_bengali(text):
    fixes = {
        "রজ্য": "রাজ্য",
        "মমল": "মামলা",
        "ববরণ": "বিবরণ",
        "পশ্চম": "পশ্চিম",
        "ইজর": "ইজারা",
        "জনাব" : "শ্রী"
    }
    for wrong, correct in fixes.items():
        text = text.replace(wrong, correct)
    return text

# ---------------------------
# LEGAL LOCKS
# ---------------------------
LEGAL_LOCKS_POST = {
    "mining lease": "খনি ইজারা",
    "grant order": "মঞ্জুরি আদেশ",
    "respondent": "প্রতিপক্ষ",
    "appellant": "আপিলকারী"
}

def apply_legal_locks_post(text):
    for eng, ben in LEGAL_LOCKS_POST.items():
        text = re.sub(re.escape(eng), ben, text, flags=re.IGNORECASE)
    return text

# ---------------------------
# MAIN
# ---------------------------
if __name__ == "__main__":

    PDF_PATH = r"D:\HC\translator\eng2.pdf"
    GLOSSARY_CSV = "dict.csv"

    txt_file = open("chunk_debug.txt", "w", encoding="utf-8")
    chunk_counter = 1

    DEVICE = "cuda" if torch.cuda.is_available() else "cpu"

    model_name = "ai4bharat/indictrans2-en-indic-1B"

    tokenizer = AutoTokenizer.from_pretrained(
        model_name,
        trust_remote_code=True,
        cache_dir=os.environ['HF_HOME'],
        local_files_only=True
    )

    model = AutoModelForSeq2SeqLM.from_pretrained(
        model_name,
        trust_remote_code=True,
        torch_dtype=torch.float16 if DEVICE == "cuda" else torch.float32,
        cache_dir=os.environ['HF_HOME'],
        local_files_only=True
    ).to(DEVICE)

    ip = IndicProcessor(inference=True)
    glossary = load_glossary(GLOSSARY_CSV)

    pdf = fitz.open(PDF_PATH)
    doc = Document()

    for page_num in range(len(pdf)):

        print(f"[PAGE {page_num+1}/{len(pdf)}]")

        if page_num > 0:
            doc.add_page_break()

        raw = pdf.load_page(page_num).get_text("text")

        # 🔥 NEW PIPELINE
        lines = clean_page_lines(raw)
        paragraphs = group_paragraphs(lines)

        translated_paragraphs = []

        for para in paragraphs:

            chunks = split_into_chunks(para)
            translated_chunks = []

            for chunk in chunks:

                txt_file.write(f"ENG_c{chunk_counter}: {chunk}\n")

                batch = ip.preprocess_batch(
                    [chunk],
                    src_lang="eng_Latn",
                    tgt_lang="ben_Beng"
                )

                inputs = tokenizer(
                    batch,
                    truncation=True,
                    max_length=512,
                    padding="longest",
                    return_tensors="pt"
                ).to(DEVICE)

                with torch.no_grad():
                    tokens = model.generate(
                        **inputs,
                        max_length=512,
                        num_beams=5,
                        repetition_penalty=1.15,
                        no_repeat_ngram_size=3
                    )

                decoded = tokenizer.batch_decode(tokens, skip_special_tokens=True)
                output = ip.postprocess_batch(decoded, lang="ben_Beng")[0]

                # CLEAN PIPELINE
                output = clean_output(output)
                output = normalize_bengali(output)
                output = apply_legal_locks_post(output)
                output = apply_glossary(output, glossary)

                txt_file.write(f"BEN_c{chunk_counter}: {output}\n------\n")

                translated_chunks.append(output)
                chunk_counter += 1

            full_para = " ".join(translated_chunks)
            translated_paragraphs.append(full_para)

        # WRITE CLEAN STRUCTURE
        for para in translated_paragraphs:
            p = doc.add_paragraph()

            if is_title(para):
                run = p.add_run(para)
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                run = p.add_run(para)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            run.font.size = Pt(12)

    txt_file.close()
    doc.save("output_2.docx")

    print("✅ DONE — Structured, clean output with paragraphs.")