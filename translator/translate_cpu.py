import os
import csv
import torch
import re
import fitz
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from transformers import AutoModelForSeq2SeqLM, AutoTokenizer
from IndicTransToolkit.processor import IndicProcessor
torch.set_num_threads(os.cpu_count())
torch.set_num_interop_threads(os.cpu_count())
os.environ['HF_HOME'] = r"D:\HC\hf_cache"
MODEL_NAME = "ai4bharat/indictrans2-en-indic-dist-200M"
MAX_TOKENS = 350
BATCH_SIZE = 8
def load_glossary(csv_path):
    glossary = {}
    if not os.path.exists(csv_path):
        return glossary

    with open(csv_path, encoding='utf-8-sig') as f:
        reader = csv.reader(f)
        next(reader, None)
        for row in reader:
            if len(row) >= 2:
                glossary[row[0].strip()] = row[1].strip()

    return glossary

def apply_glossary(text, glossary):
    for eng, ben in glossary.items():
        text = text.replace(ben, ben)
    return text


def clean_page_lines(text):
    lines = text.split('\n')
    cleaned = []

    for line in lines:
        line = line.strip()

        if not line:
            cleaned.append("")
            continue

        if re.search(r'IN THE HIGH COURT|SUPREME COURT|REPORTS', line, re.I):
            continue

        if re.match(r'^\d+$', line):
            continue

        if line in list("ABCDEFGH"):
            continue

        line = re.sub(r'<.*?>', '', line)

        cleaned.append(line)

    return cleaned

def group_paragraphs(lines):
    paragraphs = []
    current = ""

    for line in lines:
        if line == "":
            if current:
                paragraphs.append(current.strip())
                current = ""
        else:
            if current.endswith(('.', ':', '?')):
                paragraphs.append(current.strip())
                current = line
            else:
                current += " " + line

    if current:
        paragraphs.append(current.strip())

    return paragraphs

def is_title(text):
    return (
        len(text.split()) <= 12 and
        (text.isupper() or ':' in text or 'Vs.' in text)
    )

def split_into_chunks(text, max_tokens=MAX_TOKENS):
    sentences = re.split(r'(?<=[.!?])\s+', text)

    chunks = []
    current = ""

    for s in sentences:
        if len((current + " " + s).split()) < max_tokens:
            current += " " + s
        else:
            chunks.append(current.strip())
            current = s

    if current:
        chunks.append(current.strip())

    return chunks

def clean_output(text):
    text = re.sub(r'\b(\w+)( \1\b)+', r'\1', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def normalize_bengali(text):
    fixes = {
        "রজ্য": "রাজ্য",
        "মমল": "মামলা",
        "ববরণ": "বিবরণ",
        "পশ্চম": "পশ্চিম",
        "ইজর": "ইজারা",
        "জনাব" : "শ্রী"
    }
    for wrong, correct in fixes.items():
        text = text.replace(wrong, correct)
    return text

LEGAL_LOCKS_POST = {
    "mining lease": "খনি ইজারা",
    "grant order": "মঞ্জুরি আদেশ",
    "respondent": "প্রতিপক্ষ",
    "appellant": "আপিলকারী"
}

def apply_legal_locks_post(text):
    for eng, ben in LEGAL_LOCKS_POST.items():
        text = re.sub(re.escape(eng), ben, text, flags=re.IGNORECASE)
    return text

def translate_batch(chunks, tokenizer, model, ip, device):

    batch = ip.preprocess_batch(
        chunks,
        src_lang="eng_Latn",
        tgt_lang="ben_Beng"
    )

    inputs = tokenizer(
        batch,
        truncation=True,
        max_length=512,
        padding=True,
        return_tensors="pt"
    ).to(device)

    with torch.inference_mode():
        tokens = model.generate(
            **inputs,
            max_length=512,
            num_beams=1,
            do_sample=False
        )

    decoded = tokenizer.batch_decode(tokens, skip_special_tokens=True)
    outputs = ip.postprocess_batch(decoded, lang="ben_Beng")

    return outputs

if __name__ == "__main__":

    PDF_PATH = r"D:\HC\translator\eng4.pdf"
    GLOSSARY_CSV = "dict.csv"

    DEVICE = "cuda" if torch.cuda.is_available() else "cpu"

    tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME, trust_remote_code=True)
    model = AutoModelForSeq2SeqLM.from_pretrained(
        MODEL_NAME,
        trust_remote_code=True
    ).to(DEVICE)

    ip = IndicProcessor(inference=True)
    glossary = load_glossary(GLOSSARY_CSV)

    pdf = fitz.open(PDF_PATH)
    doc = Document()

    for page_num in range(len(pdf)):

        print(f"[PAGE {page_num+1}/{len(pdf)}]")

        if page_num > 0:
            doc.add_page_break()

        raw = pdf.load_page(page_num).get_text("text")

        lines = clean_page_lines(raw)
        paragraphs = group_paragraphs(lines)

        translated_paragraphs = []

        for para in paragraphs:

            chunks = split_into_chunks(para)

            outputs = []
            for i in range(0, len(chunks), BATCH_SIZE):
                batch_chunks = chunks[i:i+BATCH_SIZE]
                outputs.extend(
                    translate_batch(batch_chunks, tokenizer, model, ip, DEVICE)
                )

            cleaned_outputs = []
            for output in outputs:
                output = clean_output(output)
                output = normalize_bengali(output)
                output = apply_legal_locks_post(output)
                output = apply_glossary(output, glossary)
                cleaned_outputs.append(output)

            full_para = " ".join(cleaned_outputs)
            translated_paragraphs.append(full_para)

        for para in translated_paragraphs:
            p = doc.add_paragraph()

            if is_title(para):
                run = p.add_run(para)
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                run = p.add_run(para)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            run.font.size = Pt(12)

    doc.save("optimized_output.docx")

    print("✅ DONE — Optimized & fast.")